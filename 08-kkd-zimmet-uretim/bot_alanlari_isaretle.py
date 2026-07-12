"""
Bot Alan Personelleri Isaretle
==============================
Kok klasordeki '2026 BOT ALMIS PERSONEL LISTESI.xlsx' dosyasinda adi gecen
personellerin tum renk klasorlerindeki OLUSTURULAN/AD SOYAD.docx belgelerinde
STL-9040-S3-BOT-STARLINE iceren satirin arka planini gri yapar.
Metinler aynen kalir, sadece satirin arka plani renklendirilir.

Kullanim:
    python bot_alanlari_isaretle.py

Idempotent: zaten gri yapilmis satirlari atlar, tekrar tekrar calistirilabilir.
Eski surumden kalma 'gri yazi + ustu cizili' bicimini de temizler.
"""

from pathlib import Path
import sys

BASE_DIR = Path(__file__).parent.resolve()
BOT_LISTESI = BASE_DIR / "2026 BOT ALMIŞ PERSONEL LİSTESİ.xlsx"
BOT_TEXT = "STL-9040-S3-BOT-STARLİNE"
BG_GRI = "BFBFBF"          # Word 25% gray — okunur ama belirgin
ESKI_YAZI_GRI = (0x80, 0x80, 0x80)  # eski surumden kalma yazi rengi (temizlenecek)


def tr_upper(s):
    """Turkce-dogru buyuk harf donusumu (i -> I, i -> I)."""
    return s.translate(str.maketrans('iı', 'İI')).upper()


def normalize(name):
    """Isimleri eslestirmek icin tek formata indirge: bosluk topla + TR upper."""
    return tr_upper(" ".join(str(name).split())).strip()


def read_bot_listesi():
    """Boot Excel'inden normalize edilmis isim seti dondur."""
    import openpyxl
    wb = openpyxl.load_workbook(BOT_LISTESI, data_only=True)
    ws = wb.active
    names = set()
    for i, row in enumerate(ws.iter_rows(values_only=True)):
        if i == 0:
            continue
        if row and row[0]:
            names.add(normalize(row[0]))
    return names


def find_bot_row(doc):
    """Belgedeki BOT_TEXT iceren tablo satirini bul."""
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if BOT_TEXT in cell.text:
                    return row
    return None


def get_cell_shading(cell):
    """Hucrenin mevcut arka plan rengini dondur (hex, yoksa None)."""
    from docx.oxml.ns import qn
    tc_pr = cell._tc.tcPr
    if tc_pr is None:
        return None
    shd = tc_pr.find(qn('w:shd'))
    if shd is None:
        return None
    fill = shd.get(qn('w:fill'))
    return fill.upper() if fill else None


def set_cell_shading(cell, color_hex):
    """Hucrenin arka planini color_hex yap (var olan shd'yi silip yenisini koyar)."""
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc_pr = cell._tc.get_or_add_tcPr()
    for old in tc_pr.findall(qn('w:shd')):
        tc_pr.remove(old)
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), color_hex)
    tc_pr.append(shd)


def has_old_text_marking(cell):
    """Eski 'gri yazi + ustu cizili' bicimi var mi?"""
    from docx.shared import RGBColor
    eski_gri = RGBColor(*ESKI_YAZI_GRI)
    for para in cell.paragraphs:
        for run in para.runs:
            if run.font.strike:
                return True
            if run.font.color is not None and run.font.color.rgb == eski_gri:
                return True
    return False


def cleanup_old_text_marking(cell):
    """Eski gri+strike bicimini temizle, yazi normale donsun."""
    from docx.shared import RGBColor
    eski_gri = RGBColor(*ESKI_YAZI_GRI)
    for para in cell.paragraphs:
        for run in para.runs:
            if run.font.color is not None and run.font.color.rgb == eski_gri:
                run.font.color.rgb = None
            if run.font.strike:
                run.font.strike = None


def row_needs_update(row):
    """Satir guncellenmesi gerekiyorsa True. Tum hucreler dogru renkte
    bicimde ve eski yazi bicimi yoksa False."""
    for cell in row.cells:
        if get_cell_shading(cell) != BG_GRI:
            return True
        if has_old_text_marking(cell):
            return True
    return False


def mark_row(row):
    """Satirdaki tum hucrelerin arka planini gri yap, eski bicimi temizle."""
    for cell in row.cells:
        cleanup_old_text_marking(cell)
        set_cell_shading(cell, BG_GRI)


def process_folder(klasor, bot_set):
    """Bir renk klasorunun OLUSTURULAN'ini isle.
    Returns: (yeni_isaretlendi, zaten_isaretli, listede_degil, matched_set)."""
    from docx import Document

    olusturulan = klasor / "OLUSTURULAN"
    if not olusturulan.is_dir():
        return 0, 0, 0, set()

    yeni = 0
    zaten = 0
    listede_degil = 0
    matched = set()

    for docx_path in sorted(olusturulan.glob("*.docx")):
        if docx_path.name.startswith("~$"):
            continue
        ad_norm = normalize(docx_path.stem)
        if ad_norm not in bot_set:
            listede_degil += 1
            continue

        matched.add(ad_norm)
        try:
            doc = Document(str(docx_path))
            row = find_bot_row(doc)
            if row is None:
                print(f"  UYARI {docx_path.name}: BOT satiri bulunamadi")
                continue
            if not row_needs_update(row):
                zaten += 1
                continue
            mark_row(row)
            doc.save(str(docx_path))
            print(f"  OK    {docx_path.name}")
            yeni += 1
        except Exception as e:
            print(f"  HATA  {docx_path.name}: {e}")

    return yeni, zaten, listede_degil, matched


def main():
    if not BOT_LISTESI.is_file():
        sys.exit(f"HATA: Boot listesi bulunamadi: {BOT_LISTESI}")

    bot_set = read_bot_listesi()
    print(f"Boot listesi: {len(bot_set)} kisi (kok klasordeki Excel'den)")
    print()

    renk_klasorleri = [
        d for d in BASE_DIR.iterdir()
        if d.is_dir() and (d / "OLUSTURULAN").is_dir()
    ]
    if not renk_klasorleri:
        sys.exit("HATA: OLUSTURULAN iceren renk klasoru bulunamadi")

    toplam_yeni = toplam_zaten = toplam_listede_degil = 0
    all_matched = set()

    for klasor in sorted(renk_klasorleri):
        print(f"--- {klasor.name} ---")
        yeni, zaten, listede_degil, matched = process_folder(klasor, bot_set)
        all_matched |= matched
        toplam_yeni += yeni
        toplam_zaten += zaten
        toplam_listede_degil += listede_degil
        print(f"  yeni: {yeni}, zaten gri: {zaten}, listede degil: {listede_degil}")
        print()

    print("=" * 50)
    print(f"TOPLAM yeni isaretlendi : {toplam_yeni}")
    print(f"TOPLAM zaten griydi     : {toplam_zaten}")
    print(f"TOPLAM listede degil    : {toplam_listede_degil}")
    print()

    eslesmeyen = sorted(bot_set - all_matched)
    eslesmeyen_dosya = BASE_DIR / "eslesmeyenler.txt"
    if eslesmeyen:
        print(f"Boot listesinde olup hicbir renk klasorunde docx'i olmayan {len(eslesmeyen)} kisi:")
        for ad in eslesmeyen:
            print(f"  - {ad}")
        eslesmeyen_dosya.write_text(
            "\n".join(eslesmeyen) + "\n",
            encoding="utf-8",
        )
        print()
        print(f"({len(eslesmeyen)} isim {eslesmeyen_dosya.name} dosyasina da yazildi.)")
        print("(Excel'de yazim hatasi olabilir, kontrol et.)")
    else:
        if eslesmeyen_dosya.exists():
            eslesmeyen_dosya.unlink()
        print("Boot listesindeki herkes en az bir renk klasorunde bulundu.")


if __name__ == "__main__":
    main()
